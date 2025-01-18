#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
نظام توليد دراسات الجدوى - 2025
نظام متكامل لتوليد دراسات جدوى تفصيلية للمشاريع المختلفة في السوق المصري
جميع الحقوق محفوظة لشركة Green Light © 2025
"""

import os
import json
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
import google.generativeai as genai

class FeasibilityStudyGenerator:
    def __init__(self):
        """
        تهيئة مولد دراسات الجدوى للمشاريع المصرية.
        
        يقوم هذا المُنشئ بتهيئة نموذج Google Gemini وإعداد هيكل أقسام دراسة الجدوى. يتضمن التهيئة:
        - التحقق من وجود مفتاح API
        - تكوين نموذج Gemini للذكاء الاصطناعي
        - إنشاء هيكل منظم لأقسام دراسة الجدوى يشمل 8 أقسام رئيسية
        
        يرفع استثناء ValueError إذا لم يتم العثور على مفتاح API.
        
        الأقسام تغطي جوانب متعددة تشمل:
        - الجدوى الاقتصادية
        - الجدوى الفنية
        - الجدوى التسويقية
        - الجدوى القانونية
        - الجدوى الاجتماعية والبيئية
        - إدارة المخاطر
        - الخطط التوسعية
        - الخاتمة
        """
        # تحميل مفتاح API
        api_key = "AIzaSyCV9Xr7syuMEeXW7-H9Favc4er7GORNgxM"
        
        if not api_key:
            print("لم يتم العثور على مفتاح API")
            raise ValueError("لم يتم العثور على مفتاح API. يرجى التأكد من وجود مفتاح API")
        
        # تهيئة Google Gemini
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-pro')
        
        # هيكل الأقسام
        self.sections = {
            "مقدمة": ["أهمية دراسات الجدوى وفوائدها"],
            "القسم الأول: دراسة الجدوى الاقتصادية": [
                "المقدمة", "الهدف", "رأس المال المطلوب", "التكاليف الاستثمارية الأولية",
                "التكاليف التشغيلية السنوية", "الإيرادات السنوية", "حساب الإيرادات اليومية والسنوية",
                "إجمالي التكاليف السنوية", "التحليل المالي", "صافي الربح"
            ],
            "القسم الثاني: دراسة الجدوى الفنية": [
                "الهدف", "اختيار الموقع", "تقدير استهلاك العلف السنوي", 
                "المعدات اللازمة", "العمالة اللازمة"
            ],
            "القسم الثالث: دراسة الجدوى التسويقية": [
                "الهدف", "تحليل السوق المستهدف", "الشرائح المستهدفة",
                "تحليل المنافسة", "استراتيجيات التسويق", "التوزيع",
                "التسعير", "الترويج"
            ],
            "القسم الرابع: دراسة الجدوى القانونية": [
                "الهدف", "التراخيص والتصاريح اللازمة", "قوانين العمل والعمالة",
                "اللوائح الصحية", "القوانين البيئية"
            ],
            "القسم الخامس: دراسة الجدوى الاجتماعية والبيئية": [
                "الهدف", "التأثير الاجتماعي", "خلق فرص العمل",
                "تحسين مستوى معيشة المجتمع", "التأثير على الصحة العامة",
                "التأثير البيئي", "التأثير على الموارد الطبيعية",
                "إدارة المخلفات البيئية", "الانبعاثات الكربونية",
                "الاستدامة البيئية والاجتماعية", "الاستدامة الاجتماعية",
                "الاستدامة البيئية", "الفوائد المستقبلية"
            ],
            "القسم السادس: المخاطر المحتملة وكيفية التعامل معها": [
                "المخاطر المالية", "كيفية التعامل مع المخاطر المالية",
                "المخاطر الفنية", "المخاطر القانونية",
                "خطط واضحة للتقليل من التأثير السلبي"
            ],
            "القسم السابع: الخطة التوسعية والتوقعات المستقبلية": [
                "الخطة التوسعية", "تنويع المنتجات", "التوسع في الأسواق",
                "استخدام التكنولوجيا الحديثة", "التوقعات المستقبلية للمشروع"
            ],
            "القسم الثامن: الخاتمة": [""]
        }

    def generate_content(self, project_name, section):
        """
        توليد محتوى لقسم محدد من دراسة الجدوى باستخدام نموذج الذكاء الاصطناعي.
        
        Parameters:
            project_name (str): اسم المشروع المراد إنشاء دراسة الجدوى له
            section (str): القسم المحدد من دراسة الجدوى المطلوب توليد محتواه
        
        Returns:
            str: المحتوى المولد للقسم المحدد، أو سلسلة فارغة في حالة فشل التوليد
        
        Raises:
            يتم معالجة الأخطاء داخلياً ويتم طباعة رسائل الخطأ
        
        Notes:
            - يستخدم تأخير زمني لتجنب تجاوز حدود واجهة برمجة التطبيقات
            - يدعم توليد محتوى باللغة العربية لأقسام مختلفة من دراسة الجدوى
            - يعتمد على قاموس محدد مسبقاً للأقسام والنماذج الخاصة بكل قسم
        """
        current_year = "2025"
        prompts = {
            "مقدمة": f"""اكتب مقدمة شاملة لدراسة جدوى {project_name} في مصر لعام {current_year}.
            يجب أن تتضمن المقدمة:
            - أهمية المشروع في السوق المصري
            - الفرص المتاحة في السوق
            - لماذا يعتبر هذا المشروع مناسباً للسوق المصري
            - التحديات والفرص في السوق المصري
            اكتب المحتوى باللغة العربية، واستخدم العملة المصرية (الجنيه المصري) في جميع التكاليف.""",
            
            "القسم الأول: دراسة الجدوى الاقتصادية": f"""اكتب دراسة الجدوى الاقتصادية التفصيلية ل{project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - التكاليف الاستثمارية بالجنيه المصري
            - تكاليف التشغيل الشهرية
            - الإيرادات المتوقعة
            - فترة استرداد رأس المال
            - معدل العائد على الاستثمار
            - تحليل نقطة التعادل
            استخدم أسعار وتكاليف واقعية للسوق المصري في {current_year}.""",
            
            "القسم الثاني: دراسة الجدوى الفنية": f"""اكتب دراسة الجدوى الفنية التفصيلية ل{project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - المتطلبات الفنية للمشروع
            - المعدات والآلات المطلوبة مع أسعارها بالجنيه المصري
            - المساحة المطلوبة والموقع المناسب
            - العمالة المطلوبة ومؤهلاتهم
            - خطوات الإنتاج أو تقديم الخدمة
            استخدم معلومات تقنية حديثة ومناسبة للسوق المصري.""",
            
            "القسم الثالث: دراسة الجدوى التسويقية": f"""اكتب دراسة الجدوى التسويقية التفصيلية ل{project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - تحليل السوق المصري
            - المنافسين في السوق
            - الفئة المستهدفة
            - استراتيجيات التسويق المناسبة للسوق المصري
            - قنوات التوزيع
            - تحليل الأسعار في السوق المصري""",
            
            "القسم الرابع: دراسة الجدوى القانونية": f"""اكتب دراسة الجدوى القانونية التفصيلية ل{project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - الشكل القانوني للمشروع
            - التراخيص المطلوبة
            - الإجراءات القانونية
            - التكاليف القانونية بالجنيه المصري
            - المتطلبات الحكومية
            استخدم المعلومات القانونية الحديثة المطبقة في مصر.""",
            
            "القسم الخامس: دراسة الجدوى الاجتماعية والبيئية": f"""اكتب دراسة الجدوى الاجتماعية والبيئية ل{project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - الأثر الاجتماعي للمشروع
            - فرص العمل التي سيوفرها
            - الآثار البيئية
            - إجراءات الحماية البيئية
            - المسؤولية الاجتماعية""",
            
            "القسم السادس: المخاطر المحتملة وكيفية التعامل معها": f"""اكتب تحليلاً للمخاطر المحتملة ل{project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - المخاطر السوقية
            - المخاطر المالية
            - المخاطر التشغيلية
            - المخاطر القانونية
            - استراتيجيات إدارة المخاطر""",
            
            "القسم السابع: الخطة التوسعية والتوقعات المستقبلية": f"""اكتب الخطة التوسعية والتوقعات المستقبلية ل{project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - خطط التوسع المستقبلية
            - التوقعات المالية للسنوات القادمة
            - فرص النمو في السوق المصري
            - التطورات المتوقعة في المجال
            استخدم توقعات واقعية تناسب السوق المصري.""",
            
            "القسم الثامن: الخاتمة": f"""اكتب خاتمة لدراسة جدوى {project_name} في مصر لعام {current_year}.
            يجب أن تتضمن:
            - ملخص لأهم النقاط في الدراسة
            - توصيات نهائية
            - نظرة مستقبلية للمشروع"""
        }
        
        try:
            print(f"جاري توليد المحتوى لـ {project_name} - {section}")
            
            prompt = prompts.get(section, "")
            if not prompt:
                print(f"لم يتم العثور على قسم {section} في قوائم الأقسام")
                return ""
            
            time.sleep(2)  # تأخير لتجنب تجاوز حدود API
            response = self.model.generate_content(prompt)
            if response and response.text:
                print(f"تم توليد المحتوى بنجاح لـ {section}")
                return response.text
            else:
                print(f"فشل في توليد المحتوى لـ {section}: لا يوجد نص في الاستجابة")
                return ""
                
        except Exception as e:
            print(f"خطأ في توليد محتوى القسم {section}: {str(e)}")
            return ""

    def create_feasibility_study(self, project_name, output_dir):
        """
        Create a comprehensive feasibility study document for a specific project.
        
        This method generates a detailed Word document containing a feasibility study for the given project. The document includes:
        - A title page with the project name
        - Important notes about price estimates and market context
        - Table of contents
        - Detailed sections covering various aspects of project feasibility
        - A copyright footer
        
        Parameters:
            project_name (str): The name of the project for which the feasibility study is being generated
            output_dir (str): The directory where the generated feasibility study document will be saved
        
        Raises:
            Exception: If there are any errors during document creation or content generation
        
        Notes:
            - Uses python-docx library for document creation
            - Generates content dynamically using self.generate_content method
            - Sets specific document formatting including margins and page breaks
            - Saves the document as a .docx file in the specified output directory
        """
        try:
            doc = Document()
            
            # إعداد الصفحة
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # صفحة العنوان
            title = doc.add_heading(f"دراسة جدوى {project_name}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # إضافة معلومات هامة
            doc.add_paragraph()
            important_note = doc.add_paragraph("هام جدا:")
            important_note.runs[0].bold = True
            doc.add_paragraph("الأسعار المذكورة في هذه الدراسة هي أسعار تقديرية تستند إلى السوق المصري وتواكب المتغيرات الاقتصادية الحالية لعام 2024")
            doc.add_paragraph("يمكن تكييف دراسة الجدوى وفقًا لرأس المال المتاح، بما يلبي احتياجاتك ويحقق أهدافك الاستثمارية.")
            
            doc.add_page_break()
            
            # جدول المحتويات
            doc.add_heading("جدول المحتويات", level=1)
            for section_name, subsections in self.sections.items():
                p = doc.add_paragraph()
                p.add_run(section_name).bold = True
                for subsection in subsections:
                    if subsection:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Pt(20)
                        p.add_run(subsection)
            
            doc.add_page_break()
            
            # إنشاء محتوى الدراسة
            for section_name, subsections in self.sections.items():
                # إضافة عنوان القسم
                doc.add_heading(section_name, level=1)
                
                # إضافة المحتوى لكل قسم فرعي
                for subsection in subsections:
                    if subsection:
                        doc.add_heading(subsection, level=2)
                        content = self.generate_content(project_name, section_name)
                        doc.add_paragraph(content)
                
                # إضافة فاصل صفحات بين الأقسام الرئيسية
                doc.add_page_break()
            
            # إضافة حقوق الملكية
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Green Light © 2025 - جميع الحقوق محفوظة"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # حفظ الملف
            output_file = os.path.join(output_dir, f"{project_name}.docx")
            doc.save(output_file)
            print(f"تم إنشاء دراسة الجدوى لـ {project_name} بنجاح")
            
        except Exception as e:
            print(f"خطأ في إنشاء دراسة الجدوى لـ {project_name}: {str(e)}")

    def generate_all_studies(self):
        """
        يولد دراسات الجدوى لمجموعة متنوعة من المشاريع في السوق المصري.
        
        يقوم بالعمليات التالية:
        - قراءة قائمة المشاريع من ملف JSON
        - إنشاء المجلدات اللازمة للتصنيفات المختلفة للمشاريع
        - توليد دراسات جدوى للفئات التالية:
          * المشاريع الصناعية
          * المشاريع الزراعية
          * المشاريع الخدمية
          * المشاريع متعددة الخدمات
        
        يعالج الأخطاء المحتملة ويوفر معلومات تفصيلية أثناء عملية التوليد.
        
        ملاحظات:
        - يتطلب وجود ملف projects.json
        - يقوم بإنشاء مجلد generated_studies تلقائيًا
        - يطبع تفاصيل التقدم والأخطاء أثناء التنفيذ
        
        استثناءات:
        - يتعامل مع أي أخطاء قد تحدث أثناء القراءة أو التوليد
        """
        try:
            print("بدء عملية توليد دراسات الجدوى...")
            
            # قراءة قائمة المشاريع من الملف
            print("جاري قراءة ملف المشاريع...")
            with open('projects.json', 'r', encoding='utf-8') as f:
                projects = json.load(f)
            
            # إنشاء المجلدات إذا لم تكن موجودة
            print("جاري إنشاء المجلدات...")
            base_dir = "generated_studies"
            industrial_dir = os.path.join(base_dir, "المشاريع_الصناعية")
            agricultural_dir = os.path.join(base_dir, "المشاريع_الزراعية")
            service_dir = os.path.join(base_dir, "المشاريع_الخدمية")
            multi_service_dir = os.path.join(base_dir, "المشاريع_متعددة_الخدمات")
            
            os.makedirs(industrial_dir, exist_ok=True)
            os.makedirs(agricultural_dir, exist_ok=True)
            os.makedirs(service_dir, exist_ok=True)
            os.makedirs(multi_service_dir, exist_ok=True)
            print(f"تم إنشاء المجلدات في {os.path.abspath(base_dir)}")
            
            # توليد دراسات للمشاريع الصناعية
            print("\nجاري توليد دراسات الجدوى للمشاريع الصناعية...")
            for project in projects["industrial_projects"]:
                print(f"\nبدء العمل على: {project}")
                self.create_feasibility_study(project, industrial_dir)
            
            # توليد دراسات للمشاريع الزراعية
            print("\nجاري توليد دراسات الجدوى للمشاريع الزراعية...")
            for project in projects["agricultural_projects"]:
                print(f"\nبدء العمل على: {project}")
                self.create_feasibility_study(project, agricultural_dir)
            
            # توليد دراسات للمشاريع الخدمية
            print("\nجاري توليد دراسات الجدوى للمشاريع الخدمية...")
            for project in projects["service_projects"]:
                print(f"\nبدء العمل على: {project}")
                self.create_feasibility_study(project, service_dir)
            
            # توليد دراسات للمشاريع متعددة الخدمات
            print("\nجاري توليد دراسات الجدوى للمشاريع متعددة الخدمات...")
            for project in projects["multi_service_projects"]:
                print(f"\nبدء العمل على: {project}")
                self.create_feasibility_study(project, multi_service_dir)
            
            print("\nتم الانتهاء من توليد جميع دراسات الجدوى")
            
        except Exception as e:
            print(f"خطأ في توليد دراسات الجدوى: {str(e)}")
            import traceback
            print(traceback.format_exc())

if __name__ == "__main__":
    generator = FeasibilityStudyGenerator()
    generator.generate_all_studies()
